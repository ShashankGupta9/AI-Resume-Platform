from pathlib import Path
from typing import Union

import fitz  # PyMuPDF
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ResumeExtractionError(Exception):
    """Raised when resume text cannot be extracted."""


def extract_text_from_pdf(file_path: Union[str, Path]) -> str:
    """
    Extract text from a PDF using PyMuPDF.
    """

    file_path = Path(file_path)

    try:
        document = fitz.open(file_path)

        pages = []

        for page in document:
            text = page.get_text("text")

            if text:
                pages.append(text)

        document.close()

        extracted_text = "\n".join(pages)

        if not extracted_text.strip():
            raise ResumeExtractionError(
                "No readable text found in PDF."
            )

        return extracted_text

    except ResumeExtractionError:
        raise

    except Exception as exc:
        raise ResumeExtractionError(
            f"Failed to extract PDF text: {exc}"
        ) from exc


def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text from a DOCX using python-docx.
    """

    file_path = Path(file_path)

    try:
        document = Document(file_path)

        paragraphs = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        # Also extract text from tables.
        for table in document.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text:
                        row_text.append(cell_text)

                if row_text:
                    paragraphs.append(" | ".join(row_text))

        extracted_text = "\n".join(paragraphs)

        if not extracted_text.strip():
            raise ResumeExtractionError(
                "No readable text found in DOCX."
            )

        return extracted_text

    except ResumeExtractionError:
        raise

    except Exception as exc:
        raise ResumeExtractionError(
            f"Failed to extract DOCX text: {exc}"
        ) from exc


def extract_text(file_path: Union[str, Path]) -> str:
    """
    Automatically select the correct extractor
    based on file extension.
    """

    file_path = Path(file_path)

    extension = file_path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ResumeExtractionError(
            f"Unsupported file type: {extension}. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)

    if extension == ".docx":
        return extract_text_from_docx(file_path)

    raise ResumeExtractionError(
        f"Unsupported file type: {extension}"
    )